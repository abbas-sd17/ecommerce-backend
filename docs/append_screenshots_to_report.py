"""
Append API / admin screenshots to your project report Word file.

Default source (your filled report):
  ../Syed_Abbas_Raza_Project_Report.docx   (folder above ecommerce_backend/)

Outputs:
  ../Syed_Abbas_Raza_Project_Report_With_Screenshots.docx
  ./Syed_Abbas_Raza_Project_Report_With_Screenshots.docx   (copy under docs/ for the repo)

Optional — empty Scaler template only:
  python docs/append_screenshots_to_report.py --source template

Run from ecommerce_backend:
  pip install python-docx
  python docs/append_screenshots_to_report.py
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Inches

DOCS = Path(__file__).resolve().parent
REPO_ROOT = DOCS.parent  # ecommerce_backend
WORKSPACE_ROOT = REPO_ROOT.parent  # syedbackendprojectscaler (parent of repo)

SHOTS = DOCS / "screenshots"

SYED_SOURCE = WORKSPACE_ROOT / "Syed_Abbas_Raza_Project_Report.docx"
SYED_OUT_WS = WORKSPACE_ROOT / "Syed_Abbas_Raza_Project_Report_With_Screenshots.docx"
SYED_OUT_DOCS = DOCS / "Syed_Abbas_Raza_Project_Report_With_Screenshots.docx"

SCALER_TEMPLATE = WORKSPACE_ROOT / (
    "Scaler Neovarsity _ Academy Project Report Template (Backend Specialization).docx"
)
TEMPLATE_OUT = DOCS / "Project_Report_Template_with_Screenshots.docx"

FIGURES = [
    (
        "Figure A1 — GET /api/products/ — product list (Lecture 1: DRF & serializers).",
        SHOTS / "01-api-products-list.png",
    ),
    (
        "Figure A2 — GET /api/products/filter/ — price range filter (Lecture 2: Q objects).",
        SHOTS / "02-api-products-filter.png",
    ),
    (
        "Figure A3 — Django administration — login.",
        SHOTS / "03-django-admin-login.png",
    ),
    (
        "Figure A4 — Django administration — registered models (Products, Payments, Users, …).",
        SHOTS / "04-django-admin-index.png",
    ),
    (
        "Figure A5 — GET /api/products/1/stock/ — stock check (Lecture 4: custom handling).",
        SHOTS / "05-api-product-stock.png",
    ),
    (
        "Figure A6 — GET /api/payments/ — payment list with idempotency & webhook fields "
        "(Lectures 8–9).",
        SHOTS / "06-api-payments-list.png",
    ),
    (
        "Figure A7 — GET /api/products/health/ — health check (Lecture 10).",
        SHOTS / "07-api-products-health.png",
    ),
    (
        "Figure A8 — GET /api/products/search/ returns 405 — search endpoint is POST-only "
        "(Lecture 10).",
        SHOTS / "08-api-products-search-get-405.png",
    ),
    (
        "Figure A9 — POST /api/products/search/ — pagination & search results (Lecture 10).",
        SHOTS / "09-api-products-search-post-200.png",
    ),
    (
        "Figure A10 — GET /api/products/1/ — product detail with Redis caching "
        "(Lectures 1 & 10).",
        SHOTS / "10-api-product-detail-get.png",
    ),
]


def _append_appendix(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix — Project screenshots", level=1)
    doc.add_paragraph(
        "The following figures were captured from the local e-commerce backend "
        "(Django / Django REST Framework) at http://127.0.0.1:8000/."
    )
    for caption, path in FIGURES:
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.bold = True
        doc.add_picture(str(path), width=Inches(6.25))


def _check_figures() -> None:
    for _, p in FIGURES:
        if not p.is_file():
            raise SystemExit(f"Missing screenshot: {p}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Append screenshot appendix to Word report.")
    parser.add_argument(
        "--source",
        choices=("syed", "template"),
        default="syed",
        help="syed = Syed_Abbas_Raza_Project_Report.docx; template = empty Scaler template",
    )
    args = parser.parse_args()
    _check_figures()

    if args.source == "syed":
        if not SYED_SOURCE.is_file():
            raise SystemExit(
                f"Report not found: {SYED_SOURCE}\n"
                "Place Syed_Abbas_Raza_Project_Report.docx next to the ecommerce_backend folder."
            )
        doc = Document(str(SYED_SOURCE))
        _append_appendix(doc)
        doc.save(str(SYED_OUT_WS))
        print(f"Wrote {SYED_OUT_WS}")
        doc.save(str(SYED_OUT_DOCS))
        print(f"Wrote {SYED_OUT_DOCS}")
    else:
        if not SCALER_TEMPLATE.is_file():
            raise SystemExit(f"Template not found: {SCALER_TEMPLATE}")
        doc = Document(str(SCALER_TEMPLATE))
        _append_appendix(doc)
        doc.save(str(TEMPLATE_OUT))
        print(f"Wrote {TEMPLATE_OUT}")


if __name__ == "__main__":
    main()
